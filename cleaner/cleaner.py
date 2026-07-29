"""
==========================================================
 WhatsApp Word Cleaner Pro
 Version : 1.0.0
 Author  : Abdullah Çınar
 License : MIT
==========================================================
"""

from __future__ import annotations

from collections.abc import Iterator

from docx.document import Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

from cleaner.patterns import (
    DATE_PATTERN,
    LINE_NUMBER_PATTERN,
    MULTIPLE_SPACE_PATTERN,
    ZERO_WIDTH_PATTERN,
)

from cleaner.statistics import Statistics


class DocumentCleaner:
    """
    Main document cleaning engine.
    """

    def __init__(self):
        self.stats = Statistics()
        
        # --- Config Toggles ---
        self.normalize_paragraph_spacing = True
        self.normalize_line_spacing = True
        self.remove_extra_blank_paragraphs = True

    # ---------------------------------------------------------
    # PUBLIC
    # ---------------------------------------------------------

    def clean(self, document: Document) -> Document:
        """
        Clean complete document.
        """
        previous_blank = False

        for paragraph in self._iter_paragraphs(document):
            self.stats.paragraphs_processed += 1
            previous_blank = self._clean_paragraph(
                paragraph,
                previous_blank,
            )

        return document

    # ---------------------------------------------------------
    # ITERATORS
    # ---------------------------------------------------------

    def _iter_paragraphs(
        self,
        document: Document,
    ) -> Iterator[Paragraph]:
        # Normal paragraphs
        for paragraph in document.paragraphs:
            yield paragraph

        # Tables
        for table in document.tables:
            yield from self._iter_table(table)

        # Headers / Footers
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                yield paragraph

            for paragraph in section.footer.paragraphs:
                yield paragraph

    def _iter_table(
        self,
        table: Table,
    ) -> Iterator[Paragraph]:
        for row in table.rows:
            for cell in row.cells:
                yield from self._iter_cell(cell)

    def _iter_cell(
        self,
        cell: _Cell,
    ) -> Iterator[Paragraph]:
        for paragraph in cell.paragraphs:
            yield paragraph

        for table in cell.tables:
            yield from self._iter_table(table)

    # ---------------------------------------------------------
    # RUN ITERATOR
    # ---------------------------------------------------------

    def _iter_runs(
        self,
        paragraph: Paragraph,
    ):
        """
        Iterate over non-empty runs.
        """
        for run in paragraph.runs:
            if run.text:
                yield run

    # ---------------------------------------------------------
    # PARAGRAPH CLEANER
    # ---------------------------------------------------------

    def _clean_paragraph(
        self,
        paragraph: Paragraph,
        previous_blank: bool,
    ) -> bool:
        self._remove_dates(paragraph)
        self._remove_zero_width(paragraph)
        self._normalize_spaces(paragraph)
        
        # Paragraf ve satır aralığı boşluklarını normalleştir
        if self.normalize_paragraph_spacing or self.normalize_line_spacing:
            self._normalize_spacing(paragraph)

        text = paragraph.text.strip()

        if LINE_NUMBER_PATTERN.fullmatch(text):
            self._delete_paragraph(paragraph)
            self.stats.numbers_removed += 1
            return True

        if not text:
            if previous_blank and self.remove_extra_blank_paragraphs:
                self._delete_paragraph(paragraph)
                self.stats.blank_lines_removed += 1
            return True

        return False

    # ---------------------------------------------------------
    # CLEANING METHODS
    # ---------------------------------------------------------

    def _remove_dates(
        self,
        paragraph: Paragraph,
    ) -> None:
        """
        Remove WhatsApp date/time patterns.
        """
        for run in paragraph.runs:
            if not run.text:
                continue

            matches = DATE_PATTERN.findall(run.text)
            if matches:
                self.stats.dates_removed += len(matches)
                run.text = DATE_PATTERN.sub(
                    "",
                    run.text,
                )

    def _remove_zero_width(
        self,
        paragraph: Paragraph,
    ) -> None:
        """
        Remove invisible unicode characters.
        """
        for run in paragraph.runs:
            if not run.text:
                continue

            matches = ZERO_WIDTH_PATTERN.findall(run.text)
            if matches:
                self.stats.zero_width_removed += len(matches)
                run.text = ZERO_WIDTH_PATTERN.sub(
                    "",
                    run.text,
                )

    def _normalize_spaces(
        self,
        paragraph: Paragraph,
    ) -> None:
        """
        Replace multiple spaces with one space.
        """
        changed = False

        for run in paragraph.runs:
            if not run.text:
                continue

            new_text = MULTIPLE_SPACE_PATTERN.sub(
                " ",
                run.text,
            )
            new_text = new_text.strip()

            if new_text != run.text:
                run.text = new_text
                changed = True

        if changed:
            self.stats.spaces_fixed += 1
            
    def _normalize_spacing(
        self, 
        paragraph: Paragraph,
    ) -> None:
        """
        Normalize paragraph spacing and line spacing.
        """
        fmt = paragraph.paragraph_format
        
        if self.normalize_paragraph_spacing:
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            
        if self.normalize_line_spacing:
            fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _delete_paragraph(
        self,
        paragraph: Paragraph,
    ) -> None:
        """
        Completely remove the paragraph element from the Word document.
        """
        p = paragraph._element
        p.getparent().remove(p)

    # ---------------------------------------------------------
    # STATISTICS
    # ---------------------------------------------------------

    def reset_statistics(self) -> None:
        """
        Reset cleaning statistics.
        """
        self.stats.reset()

    def get_statistics(self) -> Statistics:
        """
        Return statistics object.
        """
        return self.stats

    @property
    def statistics(self) -> Statistics:
        """
        Statistics property.
        """
        return self.stats

    # ---------------------------------------------------------
    # REPORT
    # ---------------------------------------------------------

    def print_summary(self) -> None:
        """
        Print cleaning summary.
        """
        print(self.stats)
