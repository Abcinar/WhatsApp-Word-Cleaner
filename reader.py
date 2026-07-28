"""
WhatsApp Word Cleaner Pro
reader.py
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterator, List

from docx import Document
from docx.document import Document as DocumentType
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


SUPPORTED_EXTENSIONS = {".docx"}


class WordReaderError(Exception):
    """Base exception for WordReader."""


class InvalidDocumentError(WordReaderError):
    """Raised when an invalid document is supplied."""


class UnsupportedFileError(WordReaderError):
    """Raised when file extension is unsupported."""


@dataclass(slots=True)
class ParagraphInfo:
    index: int
    text: str
    paragraph: Paragraph


@dataclass(slots=True)
class TableCellInfo:
    table_index: int
    row_index: int
    column_index: int
    paragraph_index: int
    text: str
    paragraph: Paragraph


class WordReader:
    """
    Production-ready .docx reader.

    Features
    --------
    - Validates file
    - Loads Word document
    - Enumerates body paragraphs
    - Enumerates headers
    - Enumerates footers
    - Enumerates table paragraphs
    """

    def __init__(self, file_path: str | Path):
        self.path = Path(file_path).expanduser().resolve()

        self._validate()

        try:
            self.document: DocumentType = Document(str(self.path))
        except Exception as exc:
            raise InvalidDocumentError(
                f"Unable to open Word document: {self.path}"
            ) from exc

    # ------------------------------------------------------------------
    # Validation
    # ------------------------------------------------------------------

    def _validate(self) -> None:
        if not self.path.exists():
            raise FileNotFoundError(self.path)

        if not self.path.is_file():
            raise InvalidDocumentError(self.path)

        if self.path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFileError(self.path.suffix)

    # ------------------------------------------------------------------
    # General
    # ------------------------------------------------------------------

    @property
    def filename(self) -> str:
        return self.path.name

    @property
    def paragraph_count(self) -> int:
        return len(self.document.paragraphs)

    @property
    def table_count(self) -> int:
        return len(self.document.tables)

    @property
    def section_count(self) -> int:
        return len(self.document.sections)

    # ------------------------------------------------------------------
    # Body
    # ------------------------------------------------------------------

    def iter_body_paragraphs(self) -> Iterator[ParagraphInfo]:
        for idx, paragraph in enumerate(self.document.paragraphs):
            yield ParagraphInfo(
                index=idx,
                text=paragraph.text,
                paragraph=paragraph,
            )

    # ------------------------------------------------------------------
    # Headers
    # ------------------------------------------------------------------

    def iter_header_paragraphs(self) -> Iterator[ParagraphInfo]:
        index = 0

        for section in self.document.sections:
            for paragraph in section.header.paragraphs:
                yield ParagraphInfo(
                    index=index,
                    text=paragraph.text,
                    paragraph=paragraph,
                )
                index += 1

    # ------------------------------------------------------------------
    # Footers
    # ------------------------------------------------------------------

    def iter_footer_paragraphs(self) -> Iterator[ParagraphInfo]:
        index = 0

        for section in self.document.sections:
            for paragraph in section.footer.paragraphs:
                yield ParagraphInfo(
                    index=index,
                    text=paragraph.text,
                    paragraph=paragraph,
                )
                index += 1

    # ------------------------------------------------------------------
    # Tables
    # ------------------------------------------------------------------

    def iter_table_paragraphs(self) -> Iterator[TableCellInfo]:
        for table_index, table in enumerate(self.document.tables):
            yield from self._iterate_table(table, table_index)

    def _iterate_table(
        self,
        table: Table,
        table_index: int,
    ) -> Iterator[TableCellInfo]:

        for row_index, row in enumerate(table.rows):
            for column_index, cell in enumerate(row.cells):
                yield from self._iterate_cell(
                    cell,
                    table_index,
                    row_index,
                    column_index,
                )

    def _iterate_cell(
        self,
        cell: _Cell,
        table_index: int,
        row_index: int,
        column_index: int,
    ) -> Iterator[TableCellInfo]:

        for paragraph_index, paragraph in enumerate(cell.paragraphs):
            yield TableCellInfo(
                table_index=table_index,
                row_index=row_index,
                column_index=column_index,
                paragraph_index=paragraph_index,
                text=paragraph.text,
                paragraph=paragraph,
            )

    # ------------------------------------------------------------------
    # Combined
    # ------------------------------------------------------------------

    def all_paragraphs(self) -> List[Paragraph]:
        paragraphs: List[Paragraph] = []

        paragraphs.extend(self.document.paragraphs)

        for section in self.document.sections:
            paragraphs.extend(section.header.paragraphs)
            paragraphs.extend(section.footer.paragraphs)

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)

        return paragraphs

    # ------------------------------------------------------------------
    # Statistics
    # ------------------------------------------------------------------

    def statistics(self) -> dict[str, int]:
        body = self.paragraph_count

        headers = sum(
            len(section.header.paragraphs)
            for section in self.document.sections
        )

        footers = sum(
            len(section.footer.paragraphs)
            for section in self.document.sections
        )

        tables = 0

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    tables += len(cell.paragraphs)

        return {
            "body_paragraphs": body,
            "header_paragraphs": headers,
            "footer_paragraphs": footers,
            "table_paragraphs": tables,
            "tables": self.table_count,
            "sections": self.section_count,
            "total_paragraphs": body + headers + footers + tables,
        }

    # ------------------------------------------------------------------
    # Access
    # ------------------------------------------------------------------

    def get_document(self) -> DocumentType:
        return self.document
